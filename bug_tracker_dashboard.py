import io
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import streamlit as st
from docx import Document
from docx.shared import Inches

plt.rcParams['font.family'] = ['Microsoft JhengHei']  # 微軟正黑體，Windows 預設有
plt.rcParams['axes.unicode_minus'] = False

st.title("Bug Tracker Dashboard")
df = None
uploaded_file = st.file_uploader("Please upload the CSV file to analyze.", type=["csv"])

if uploaded_file is not None:
  try:
    df = pd.read_csv(uploaded_file, encoding = 'big5')
  except Exception as e:
    st.error(f"Error reading the file. Please make sure the file is encoded in Big5.：{e}")

if df is not None:
  st.header("🔧 Overview KPIs")
  #All Bug
  st.metric(label="All Bugs", value=df.shape[0])
  #All status
  status_counts = df['狀態'].value_counts().reset_index()
  status_counts.columns = ['Status', 'Count']
  st.table(status_counts)

  #MTTR
  df['建立日期']=pd.to_datetime(df['建立日期'], format='%Y-%m-%d %H:%M')
  df['結束日期']=pd.to_datetime(df['結束日期'], format='%Y-%m-%d %H:%M')
  df['處理天數']=(df['結束日期']-df['建立日期']).dt.days
  df_closed = df.dropna(subset=['結束日期'])
  MTTR = df_closed['處理天數'].mean()
  st.metric(label = "Time to Close(Days)", value = round(MTTR,2))

  #Bug Status Breakdown
  st.header("📊 Diagram Overview")
  plt.figure(figsize=(8, 6))
  sns.countplot(x='狀態', data=df)
  plt.xticks(rotation=45)
  plt.title('Current Bug Status Breakdown')
  st.pyplot(plt)
  plt.savefig('status_plot.png')
  plt.close()

  #Severity vs Priority Heatmap
  ct = pd.crosstab(df['Severity'], df['優先權'])
  plt.figure(figsize=(8,6))
  sns.heatmap(ct, annot = True, fmt = 'd', cmap='YlGnBu')
  plt.title('Severity vs Priority Heatmap')
  st.pyplot(plt)
  plt.savefig('heatmap_plot.png')
  plt.close()

  #MTTR
  plt.figure(figsize=(8, 6))
  sns.violinplot(x='優先權', y='處理天數', data=df)
  plt.title("Processing Time Distribution by Priority (Violin Plot)")
  plt.ylabel("Processing Time (days)")
  plt.xlabel("Priority")
  st.pyplot(plt)
  plt.savefig('violin_plot.png')
  plt.close()

  if st.button("Export to Word"):
        doc = Document()
        doc.add_heading('Bug Tracker Report', 0)

        # KPI Summary
        doc.add_heading('Overview KPIs', level=1)
        doc.add_paragraph(f"All Bugs: {df.shape[0]}")
        doc.add_paragraph(f"Time to Close (Days): {round(MTTR, 2)}")

        # Status Table
        doc.add_heading('Status Counts', level=2)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Status'
        hdr_cells[1].text = 'Count'
        for _, row in status_counts.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Status'])
            row_cells[1].text = str(row['Count'])

        # 用 BytesIO 把圖表存成圖片再放進 Word
        images = ['status_plot.png', 'heatmap_plot.png', 'violin_plot.png']
        for img_path in images:
            doc.add_page_break()
            doc.add_picture(img_path, width=Inches(5))
    
        # Save doc to bytes buffer
        # 儲存到BytesIO
        word_io = io.BytesIO()
        doc.save(word_io)
        word_io.seek(0)

        st.download_button(
            label="Download Word Report",
            data=word_io,
            file_name="bug_tracker_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

else:
    st.info("Please upload a CSV file to proceed.")

